from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from typing import List, Optional
from pydantic import BaseModel
from datetime import datetime

from app.db import get_db
from app.deps import get_current_active_user
from app.models import ThemeTemplate, User, Role

router = APIRouter()

# --- Schemas ---
class TemplateCreate(BaseModel):
    id: str
    name: str
    description: Optional[str] = None
    preview_url: Optional[str] = None
    actual_web_url: Optional[str] = None
    colors_json: Optional[dict] = {}
    features_json: Optional[list] = []

class TemplateResponse(BaseModel):
    id: str
    name: str
    description: Optional[str]
    preview_url: Optional[str]
    actual_web_url: Optional[str]
    colors_json: dict
    features_json: list
    created_at: datetime
    is_active: bool
    is_published: bool

    class Config:
        orm_mode = True

# --- Endpoints ---

@router.get("/templates", response_model=List[TemplateResponse])
def get_templates(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """List all active templates (Admin View)."""
    return db.query(ThemeTemplate).filter(ThemeTemplate.is_active == True).all()

@router.post("/templates/{template_id}/publish")
def toggle_template_publish(
    template_id: str,
    publish: bool,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Publish/Unpublish a template."""
    if current_user.role not in [Role.ADMIN, Role.MANAGER]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    template = db.query(ThemeTemplate).filter(ThemeTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    template.is_published = publish
    db.commit()
    return {"success": True, "message": f"Template {'published' if publish else 'unpublished'}"}

@router.post("/templates", response_model=TemplateResponse)
def create_template(
    template: TemplateCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Create a new theme template (Admin/Manager only)."""
    if current_user.role not in [Role.ADMIN, Role.MANAGER]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    existing = db.query(ThemeTemplate).filter(ThemeTemplate.id == template.id).first()
    if existing:
        raise HTTPException(status_code=400, detail="Template ID already exists")
    
    new_template = ThemeTemplate(
        id=template.id,
        name=template.name,
        description=template.description,
        preview_url=template.preview_url,
        actual_web_url=template.actual_web_url,
        colors_json=template.colors_json,
        features_json=template.features_json,
        is_published=False # Default to draft
    )
    db.add(new_template)
    db.commit()
    db.refresh(new_template)
    return new_template

@router.delete("/templates/{template_id}")
def delete_template(
    template_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Soft delete a template (Admin/Manager only)."""
    if current_user.role not in [Role.ADMIN, Role.MANAGER]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    template = db.query(ThemeTemplate).filter(ThemeTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Soft delete
    template.is_active = False
    template.is_published = False # Also unpublish
    db.commit()
    return {"success": True, "message": "Template deleted"}

# --- Bulk Upload ---
from fastapi import UploadFile, File
import pandas as pd
import io

@router.post("/templates/bulk")
async def bulk_upload_templates(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Bulk upload theme templates from Excel/CSV."""
    if current_user.role not in [Role.ADMIN, Role.MANAGER]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")

    try:
        content = await file.read()
        file_obj = io.BytesIO(content)
        
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file_obj)
        elif file.filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file_obj)
        elif file.filename.endswith('.docx'):
            import docx
            doc = docx.Document(file_obj)
            # Find the first table
            if len(doc.tables) == 0:
                raise HTTPException(status_code=400, detail="No tables found in Word document.")
            
            table = doc.tables[0]
            data = [[cell.text for cell in row.cells] for row in table.rows]
            
            if len(data) < 2:
                raise HTTPException(status_code=400, detail="Table in document is empty or has no data.")
                
            # Assume first row is header
            headers = data[0]
            rows = data[1:]
            df = pd.DataFrame(rows, columns=headers)
            
        elif file.filename.endswith('.pdf'):
            import pdfplumber
            
            data = []
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        # Filter out empty rows or None
                        cleaned_table = [[cell if cell is not None else "" for cell in row] for row in table]
                        data.extend(cleaned_table)
            
            if not data or len(data) < 2:
                raise HTTPException(status_code=400, detail="No readable tables found in PDF.")
                
            # Assume first row of the first identified table is header
            headers = data[0]
            rows = data[1:]
            df = pd.DataFrame(rows, columns=headers)
            
        else:
            raise HTTPException(status_code=400, detail="Invalid file format. Use CSV, Excel, Word (.docx), or PDF.")

        # Expected columns: "Template ID", "Template Name", "Description", "Preview Image URL", "Actual Template Web URL", "Features"
        
        # Normalize column names (strip whitespace, lower case)
        df.columns = [str(c).strip().lower().replace(' ', '_') for c in df.columns]
        print(f"Bulk Upload Columns Detected: {df.columns.tolist()}") # Debug log
        
        # Mappings helper
        def get_val(row, keys):
            for k in keys:
                if k in row:
                    val = row[k]
                    if pd.notna(val):
                        return str(val).strip()
            return ""

        success_count = 0
        errors = []

        for index, row in df.iterrows():
            try:
                # Flexible column matching
                t_id = get_val(row, ['template_id', 'id', 'templateid', 'code'])
                t_name = get_val(row, ['template_name', 'name', 'templatename', 'title'])
                
                if not t_id and t_name:
                    # Auto-generate ID from name if missing
                    # Simple slugify: lowercase, replace non-alphanumeric with hyphen
                    import re
                    t_id = re.sub(r'[^a-z0-9]+', '-', t_name.lower()).strip('-')

                if not t_id or not t_name:
                    # Log why it skipped
                    if not t_id and not t_name:
                        # Probably an empty row, ignore silently
                        continue
                    errors.append(f"Row {index+2}: Missing ID or Name (Found ID: '{t_id}', Name: '{t_name}')")
                    continue

                description = get_val(row, ['description', 'desc', 'summary'])
                preview_url = get_val(row, ['preview_image_url', 'preview_url', 'image', 'image_url', 'preview'])
                actual_web_url = get_val(row, ['actual_template_web_url', 'actual_web_url', 'web_url', 'url', 'link', 'site_url'])
                actual_web_url = get_val(row, ['actual_template_web_url', 'actual_web_url', 'web_url', 'url', 'link', 'site_url'])
                
                # Improve features detection
                features_keys = ['features', 'feature_list', 'tags', 'features_(comma_separated)']
                features_raw = get_val(row, features_keys)
                
                # Fallback: if empty, look for any column containing 'feature'
                if not features_raw:
                    for col in df.columns:
                        if 'feature' in col:
                            val = row[col]
                            if pd.notna(val):
                                features_raw = str(val).strip()
                                break

                # Check if exists
                existing = db.query(ThemeTemplate).filter(ThemeTemplate.id == t_id).first()
                
                # Parse features robustly
                import re
                features_list = []
                if features_raw:
                    features_list = [f.strip() for f in re.split(r'[,\n;]+', features_raw) if f.strip()]

                if existing:
                    # Update existing
                    existing.name = t_name
                    existing.description = description
                    existing.preview_url = preview_url
                    existing.actual_web_url = actual_web_url
                    existing.features_json = features_list
                    
                    # Re-activate if it was deleted
                    existing.is_active = True
                    
                else:
                    # Create new
                    new_template = ThemeTemplate(
                        id=t_id,
                        name=t_name,
                        description=description,
                        preview_url=preview_url,
                        actual_web_url=actual_web_url,
                        features_json=features_list,
                        colors_json={ "primary": "#000000", "secondary": "#333333", "accent": "#666666" }, # Default colors
                        is_published=False # Uploaded templates are drafts by default
                    )
                    db.add(new_template)
                
                success_count += 1
                
                success_count += 1
            except Exception as e:
                errors.append(f"Row {index+1}: {str(e)}")

        db.commit()
        return {
            "success": True, 
            "message": f"Successfully processed {success_count} templates.",
            "errors": errors
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
